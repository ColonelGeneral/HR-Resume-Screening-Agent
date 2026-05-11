from __future__ import annotations

import io
import json
from pathlib import Path
from typing import Any

import fitz
from docx import Document
from fastapi import UploadFile

from app.core.config import settings
from app.schemas import ParsedDocument
from app.utils.metadata import normalize_text

try:
    import pytesseract
    from PIL import Image
except Exception:  # pragma: no cover - optional OCR dependency
    pytesseract = None
    Image = None


class IngestionError(ValueError):
    pass


def validate_upload(filename: str, content_type: str | None = None, size_bytes: int | None = None) -> None:
    suffix = Path(filename).suffix.lower()
    if suffix not in settings.allowed_extensions:
        raise IngestionError(f"Unsupported file type: {suffix}")
    if size_bytes is not None and size_bytes > settings.max_upload_mb * 1024 * 1024:
        raise IngestionError("File exceeds max upload size")
    if content_type and content_type in {"application/x-msdownload", "application/x-sh"}:
        raise IngestionError("Executable uploads are not allowed")


def _ocr_pixmap(pixmap_bytes: bytes) -> str:
    if not settings.ocr_enabled or pytesseract is None or Image is None:
        return ""
    image = Image.open(io.BytesIO(pixmap_bytes)).convert("RGB")
    return pytesseract.image_to_string(image)


def extract_pdf_text(file_bytes: bytes, enable_ocr: bool = True) -> str:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages: list[str] = []
    for page in doc:
        text = normalize_text(page.get_text("text"))
        if not text and enable_ocr:
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            text = normalize_text(_ocr_pixmap(pix.tobytes("png")))
        pages.append(text)
    return "\n\n".join([page for page in pages if page])


def extract_docx_text(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    return "\n".join(normalize_text(para.text) for para in doc.paragraphs if normalize_text(para.text))


def extract_txt_text(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")


def extract_linkedin_json(data: dict[str, Any]) -> str:
    sections: list[str] = []
    for key in ("name", "headline", "summary", "about", "skills", "experience", "education", "projects"):
        value = data.get(key)
        if not value:
            continue
        if isinstance(value, list):
            section = ", ".join(str(item) for item in value)
        elif isinstance(value, dict):
            section = json.dumps(value, ensure_ascii=False)
        else:
            section = str(value)
        sections.append(f"{key.title()}: {section}")
    return "\n".join(sections)


async def ingest_upload_file(upload_file: UploadFile) -> ParsedDocument:
    file_bytes = await upload_file.read()
    validate_upload(upload_file.filename, upload_file.content_type, len(file_bytes))
    suffix = Path(upload_file.filename).suffix.lower()
    if suffix == ".pdf":
        text = extract_pdf_text(file_bytes)
        source_type = "pdf"
    elif suffix == ".docx":
        text = extract_docx_text(file_bytes)
        source_type = "docx"
    elif suffix == ".txt":
        text = extract_txt_text(file_bytes)
        source_type = "txt"
    elif suffix == ".json":
        raw_json = json.loads(file_bytes.decode("utf-8", errors="ignore"))
        text = extract_linkedin_json(raw_json)
        source_type = "linkedin_json"
    else:
        raise IngestionError(f"Unsupported file type: {suffix}")
    metadata = {"content_type": upload_file.content_type}
    if suffix == ".json":
        metadata["json_data"] = raw_json
    return ParsedDocument(source_name=upload_file.filename, source_type=source_type, text=normalize_text(text), metadata=metadata)
